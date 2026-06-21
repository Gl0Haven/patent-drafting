#!/usr/bin/env python
"""专利说明书附图验收器（patent-drafting）。

职责：只做静态硬检查 + 图文引用核对，不修改任何文件。
两层验收：本工具=自动硬检查；人工视觉门禁（2/3 与 4×6cm 目视文字/箭头/图例）另行执行。

用法：
    python patent_figures_check.py --config figures.json --figures-docx 说明书附图.docx [--spec 说明书.docx]

输入：图配置 JSON + 附图 DOCX + （可选）说明书 DOCX/Markdown（.doc 须先转 .docx/MD）。
仅有附图与说明书、无配置时无法判预期图名/子图/摘要，故配置必给。
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from io import BytesIO
from pathlib import Path

EMU_PER_MM = 36000


class Report:
    def __init__(self) -> None:
        self.rows: list[tuple[str, bool, str]] = []

    def add(self, name: str, ok: bool, detail: str = "") -> None:
        self.rows.append((name, ok, detail))

    def failed(self) -> int:
        return sum(1 for _, ok, _ in self.rows if not ok)

    def render(self) -> str:
        out = []
        for name, ok, detail in self.rows:
            mark = "PASS" if ok else "FAIL"
            out.append(f"[{mark}] {name}" + (f" — {detail}" if detail else ""))
        out.append("")
        out.append(f"小计：{len(self.rows)} 项，FAIL {self.failed()} 项。")
        return "\n".join(out)


def _caption_label(f: dict) -> str:
    sub = f.get("subfigure")
    return f"图{f['number']}({sub})" if sub else f"图{f['number']}"


CAP_RE = re.compile(r"^图\s*\d+\s*(\([a-zA-Z]\))?\s*$")


def check_docx(docx_path: Path, config: dict, rep: Report) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    figs = config["figures"]

    # 1. A4
    sec = doc.sections[0]
    w_mm = round(sec.page_width / EMU_PER_MM, 1)
    h_mm = round(sec.page_height / EMU_PER_MM, 1)
    a4 = abs(w_mm - 210) <= 1 and abs(h_mm - 297) <= 1
    rep.add("纸张 A4(210×297mm)", a4, f"实测 {w_mm}×{h_mm}mm")

    # 2. 页脚 PAGE 域
    foot_xml = sec.footer._element.xml if sec.footer is not None else ""
    has_page = ("PAGE" in foot_xml) and ("fldChar" in foot_xml or "fldSimple" in foot_xml)
    rep.add("页脚含 PAGE 页码域", has_page, "" if has_page else "未检出 PAGE 域")

    # 题注段（文本形如“图N(±a)”）与图片段
    cap_texts: list[str] = []
    cap_centered_all = True
    img_centered_all = True
    img_para_cnt = 0
    residual = []
    for p in doc.paragraphs:
        t = p.text.strip()
        is_drawing = p._p.findall(".//" + qn("w:drawing"))
        if is_drawing:
            img_para_cnt += 1
            if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                img_centered_all = False
        if t and CAP_RE.match(t):
            cap_texts.append(t)
            if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                cap_centered_all = False
        if t and re.search(r"```mermaid|flowchart\s+(TB|LR|TD)|graph\s+(TB|LR|TD)|-->|subgraph", t):
            residual.append(t[:40])

    # 3. 图片数 == 配置图数
    rep.add("图片数与配置一致", img_para_cnt == len(figs), f"图片段 {img_para_cnt} / 配置 {len(figs)}")
    # 4. 图片居中
    rep.add("图片居中", img_centered_all and img_para_cnt > 0, "" if img_centered_all else "存在未居中图片")
    # 5. 题注仅“图N(±a)”且居中、无长图名
    exp_caps = [_caption_label(f) for f in figs]
    caps_ok = cap_texts == exp_caps
    rep.add("题注=有序“图N(±a)”且无长图名", caps_ok, f"实测 {cap_texts}")
    rep.add("题注居中", cap_centered_all and len(cap_texts) > 0, "" if cap_centered_all else "存在未居中题注")
    # 6. 连续图号 + 子图成对
    nums = [f["number"] for f in figs]
    seq_ok = nums == sorted(set(nums)) or _continuous_with_sub(figs)
    rep.add("图号连续、子图成对", _check_numbering(figs), f"图号 {[_caption_label(f) for f in figs]}")
    # 7. 无残留 mermaid/代码
    rep.add("无残留 mermaid/流程代码", not residual, "" if not residual else f"疑似残留 {residual[:3]}")

    # 8. 图像灰度 + 有效DPI（按 docx 内嵌显示尺寸计，非裸像素——细长图不误杀）
    from PIL import Image

    gray_ok = True
    dpi_ok = True
    n_img = 0
    details = []
    for shape in doc.inline_shapes:
        try:
            blip = shape._inline.graphic.graphicData.pic.blipFill.blip
            rId = blip.get(qn("r:embed"))
            blob = doc.part.related_parts[rId].blob
        except Exception:
            continue
        n_img += 1
        with Image.open(BytesIO(blob)) as im:
            wpx, hpx = im.size
            hist = im.convert("HSV").getchannel("S").histogram()
            tot = sum(hist) or 1
            mean_s = sum(i * c for i, c in enumerate(hist)) / tot
            if mean_s > 12:
                gray_ok = False
        dw_in = (shape.width or 0) / 914400
        dh_in = (shape.height or 0) / 914400
        dpi = min(wpx / dw_in if dw_in else 0, hpx / dh_in if dh_in else 0)
        if dpi < 150:
            dpi_ok = False
        details.append(f"{wpx}x{hpx}@{dpi:.0f}dpi,S{mean_s:.1f}")
    rep.add("内嵌图为灰度/低饱和", gray_ok and n_img > 0, "; ".join(details[:6]))
    rep.add("内嵌图有效DPI≥150", dpi_ok and n_img > 0, "; ".join(details[:6]))


def _check_numbering(figs: list[dict]) -> bool:
    # 主图号应为 1..max 连续；带子图的图号其子图须成对(a,b...)且无裸主号
    from collections import defaultdict

    bynum = defaultdict(list)
    for f in figs:
        bynum[f["number"]].append(f.get("subfigure"))
    nums = sorted(bynum)
    if nums != list(range(1, max(nums) + 1)):
        return False
    for n, subs in bynum.items():
        if len(subs) == 1:
            if subs[0] is not None:
                return False  # 单图却带子图号
        else:
            letters = sorted(s for s in subs if s)
            if None in subs or letters != [chr(ord("a") + i) for i in range(len(subs))]:
                return False
    return True


def _continuous_with_sub(figs: list[dict]) -> bool:
    return _check_numbering(figs)


def _read_spec_text(spec_path: Path) -> str:
    if spec_path.suffix.lower() == ".docx":
        from docx import Document

        return "\n".join(p.text for p in Document(str(spec_path)).paragraphs)
    return spec_path.read_text(encoding="utf-8", errors="replace")


def check_spec_refs(spec_path: Path, config: dict, rep: Report) -> None:
    text = _read_spec_text(spec_path)
    figs = config["figures"]
    main_nums = sorted({f["number"] for f in figs})

    # 定位 附图说明 / 具体实施方式 区段（兼容 @H 标记或纯标题）
    def section_after(title: str) -> str:
        m = re.search(rf"(?:@H\s*)?{title}", text)
        if not m:
            return ""
        start = m.end()
        nxt = re.search(r"(?:@H\s*)?(附图说明|具体实施方式|有益效果)", text[start:])
        return text[start : start + (nxt.start() if nxt else len(text) - start)]

    fig_desc = section_after("附图说明")
    embodiment = section_after("具体实施方式")

    # a) 每个 full_name 在附图说明出现
    missing_names = [f["full_name"] for f in figs if f["full_name"] not in text]
    rep.add("各完整图名见于附图说明/正文", not missing_names,
            "" if not missing_names else f"缺 {len(missing_names)} 个，如：{missing_names[0][:24]}…")

    # b) 每个主图号在附图说明声明
    decl_missing = [n for n in main_nums if not re.search(rf"图\s*{n}(?!\d)", fig_desc)]
    rep.add("各图号在附图说明声明", not decl_missing and bool(fig_desc),
            "附图说明未找到" if not fig_desc else (f"缺图号 {decl_missing}" if decl_missing else ""))

    # c) 每个主图号在具体实施方式被引用
    ref_missing = [n for n in main_nums if not re.search(rf"图\s*{n}(?!\d)", embodiment)]
    rep.add("各图号在具体实施方式被引用", not ref_missing and bool(embodiment),
            "具体实施方式未找到" if not embodiment else (f"未引用 {ref_missing}" if ref_missing else ""))

    # d) 正文引用的图号都存在（防悬空/错引，如“由图6可见”而无图6）
    refs = {int(x) for x in re.findall(r"图\s*(\d+)", text)}
    dangling = sorted(r for r in refs if r not in main_nums)
    rep.add("正文无悬空图号引用", not dangling, "" if not dangling else f"引用了不存在的 图{dangling}")

    # e) 子图标号格式与正文一致（半角/全角统一；防附图页与正文格式分叉）
    sub_labels = [_caption_label(f) for f in figs if f.get("subfigure")]
    sub_missing = [lab for lab in sub_labels if lab not in text]
    rep.add("子图标号格式与正文一致", not sub_labels or not sub_missing,
            "" if not sub_missing else f"正文缺/格式不符：{sub_missing}")


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="专利说明书附图验收（静态硬检查 + 图文引用核对，不改文件）")
    ap.add_argument("--config", required=True, type=Path)
    ap.add_argument("--figures-docx", required=True, type=Path)
    ap.add_argument("--spec", type=Path, default=None, help="说明书 .docx 或 .md/.txt（.doc 须先转换）")
    args = ap.parse_args(argv)

    config = json.loads(args.config.read_text(encoding="utf-8"))
    rep = Report()

    try:
        check_docx(args.figures_docx.resolve(), config, rep)
    except Exception as e:
        rep.add("附图 DOCX 解析", False, str(e))

    if args.spec is not None:
        if args.spec.suffix.lower() == ".doc":
            rep.add("说明书引用核对", False, "传入的是旧版 .doc，请先转 .docx/.md 再核对")
        else:
            try:
                check_spec_refs(args.spec.resolve(), config, rep)
            except Exception as e:
                rep.add("说明书引用核对", False, str(e))
    else:
        rep.add("说明书引用核对", True, "未提供 --spec，跳过（图文双向引用未核）")

    print(rep.render())
    print("\n[人工视觉门禁] 请将各图按 2/3 缩放与 4×6cm 渲染后目视：文字/箭头/图例是否清晰、灰度是否可辨。")
    return 0 if rep.failed() == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
