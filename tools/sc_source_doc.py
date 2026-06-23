#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
sc_source_doc.py — 软著（软件著作权）源代码鉴别材料生成器（确定性、不烧 token）。

从一个或多个代码文件夹扫描源码，按 CPCC 通行口径裁剪排版成"源代码.docx"：
  - 启动/入口代码置首；
  - 去空行（可选去纯括号行/凑数注释）；
  - 每页 >=50 行（强制分页保证），不足设定页数全收，超出则取前 N 页 + 后 N 页；
  - 页眉 = 软件全称 + 版本号；右上角页码 "第 X 页 共 Y 页"；
  - 等宽字体、A4、单面。
并打印实测总代码行数（回填申请表）与裁剪明细，便于核对自洽。

用法示例：
  python sc_source_doc.py --src ./src --src ./gui --name "XX系统" --version "V1.0" --out ./源代码.docx
  python sc_source_doc.py --src ./src --name "XX系统" --version "V1.0" --dry-run

依赖：python-docx（已在 conda 环境 patent-drafting 中）。--dry-run 不需要 python-docx。
"""
import argparse
import math
import os
import sys

# 默认源码扩展名（覆盖四类常见栈：MATLAB/Qt-C++/Python/Web 等）
DEFAULT_EXTS = [
    ".m", ".mlx",                              # MATLAB
    ".c", ".cc", ".cpp", ".cxx", ".h", ".hpp", ".hxx",  # C/C++
    ".ui", ".qml",                             # Qt 界面
    ".py", ".pyw",                             # Python
    ".js", ".jsx", ".ts", ".tsx", ".vue",      # Web/Electron
    ".java", ".kt", ".cs", ".go", ".rs", ".swift", ".php", ".rb",
    ".html", ".css", ".scss",
]
# 排除目录（依赖/构建产物/版本控制/第三方）
EXCLUDE_DIRS = {
    ".git", ".svn", ".hg", "node_modules", "bower_components",
    "build", "dist", "out", "bin", "obj", "target", "release", "debug",
    "__pycache__", ".venv", "venv", "env", ".idea", ".vs", ".vscode",
    "vendor", "third_party", "thirdparty", "3rdparty", "deps", "external",
    "site-packages", "migrations", "coverage", ".pytest_cache",
}
# 测试目录（默认排除，--include-tests 可保留）：测试非核心独创代码，不宜进 60 页
TEST_DIRS = {
    "test", "tests", "__tests__", "spec", "specs", "e2e",
    "testing", "unittest", "unittests", "cypress",
}
# 入口/启动文件名优先级（小写 stem）
ENTRY_STEMS = ["__main__", "main", "app", "index", "startup", "program", "run"]


def is_generated(name):
    """常见自动生成代码（非作者独创表达），默认不计入。"""
    low = name.lower()
    if name.startswith(("moc_", "ui_", "qrc_", "rcc_")):
        return True
    gen_suffix = (
        ".min.js", ".min.css",
        "_pb2.py", "_pb2_grpc.py", ".pb.go", ".pb.cc", ".pb.h",
        ".g.dart", ".freezed.dart", ".g.cs", ".designer.cs",
        ".generated.cs", ".generated.ts", "_rc.py",
    )
    return low.endswith(gen_suffix)


def is_probably_binary(path):
    try:
        with open(path, "rb") as f:
            chunk = f.read(4096)
        return b"\x00" in chunk
    except OSError:
        return True


def collect_files(srcs, exts, exclude_extra, include_tests=False, include_generated=False):
    """递归收集源码文件，返回 (files, excluded_dirs, stats)。
    默认排除测试目录与自动生成代码（--include-tests / --include-generated 可保留）。"""
    exts = {e.lower() for e in exts}
    skip_dirs = set(EXCLUDE_DIRS) | set(exclude_extra)
    if not include_tests:
        skip_dirs |= TEST_DIRS
    excluded_dirs = set()
    files = []
    seen = set()
    n_generated = 0
    for src in srcs:
        src = os.path.abspath(src)
        if not os.path.isdir(src):
            print(f"[警告] 跳过不存在的目录: {src}", file=sys.stderr)
            continue
        for root, dirs, names in os.walk(src):
            # 原地裁剪要遍历的子目录
            kept = []
            for d in dirs:
                if d.lower() in skip_dirs:
                    excluded_dirs.add(d)
                else:
                    kept.append(d)
            dirs[:] = kept
            for name in names:
                ext = os.path.splitext(name)[1].lower()
                if ext not in exts:
                    continue
                if not include_generated and is_generated(name):
                    n_generated += 1
                    continue
                p = os.path.join(root, name)
                rp = os.path.realpath(p)
                if rp in seen:
                    continue
                if is_probably_binary(p):
                    continue
                seen.add(rp)
                files.append(p)
    return files, sorted(excluded_dirs), {"generated_skipped": n_generated}


def order_files(files, tail_pat=None):
    """入口/启动文件置首，其余按相对路径稳定排序；tail_pat 命中的文件强制置于末尾
    （便于让末页落在程序自然结尾，审核员通常会看末页是否程序结束）。"""
    tp = tail_pat.lower() if tail_pat else None

    def rank(p):
        if tp and tp in p.lower():
            return (200, p)  # 末尾
        stem = os.path.splitext(os.path.basename(p))[0].lower()
        ext = os.path.splitext(p)[1].lower()
        if ext == ".mlapp":
            return (0, p)
        if stem in ENTRY_STEMS:
            return (1 + ENTRY_STEMS.index(stem), p)
        return (100, p)
    return sorted(files, key=rank)


def line_is_droppable(line, strip_comments):
    s = line.strip()
    if s == "":
        return True  # 空行总是去掉
    if strip_comments:
        if s in ("{", "}", "(", ")", "[", "]", ";"):
            return True
        if s.startswith(("//", "#", "%", "*")) and len(s) <= 3:
            return True
    return False


def read_code_lines(files, root_for_marker, strip_comments, file_marker):
    """读取并清洗所有文件的代码行，返回 (lines, total_lines, per_file_counts, last_line)。"""
    lines = []
    per_file = []
    for p in files:
        try:
            with open(p, "r", encoding="utf-8", errors="replace") as f:
                raw = f.read().splitlines()
        except OSError as e:
            print(f"[警告] 读取失败，跳过: {p} ({e})", file=sys.stderr)
            continue
        cleaned = [ln.rstrip() for ln in raw if not line_is_droppable(ln, strip_comments)]
        if not cleaned:
            continue
        if file_marker:
            rel = os.path.relpath(p, root_for_marker) if root_for_marker else os.path.basename(p)
            lines.append(f"// ===== {rel} =====")
        lines.extend(cleaned)
        per_file.append((p, len(cleaned)))
    last_line = lines[-1] if lines else ""
    return lines, len(lines), per_file, last_line


def select_pages(lines, lpp, pages_each_side):
    """分页并按"前 N 页 + 后 N 页"裁剪，返回 (selected_lines, total_pages, truncated)。"""
    total_pages = max(1, math.ceil(len(lines) / lpp))
    limit = 2 * pages_each_side
    if total_pages <= limit:
        return lines, total_pages, False
    head = lines[: pages_each_side * lpp]
    tail = lines[-pages_each_side * lpp:]
    return head + tail, limit, True


def build_docx(selected, lpp, name, version, out_path, font, font_size):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    # A4 + 适中页边距
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    content_w = Cm(21.0 - 2.2 * 2)

    # 正文样式：等宽字体
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(font_size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    # 页眉：左 软件名+版本，右 第X页 共Y页
    header = sec.header
    hp = header.paragraphs[0]
    hp.text = f"{name} {version}".strip()
    hp.paragraph_format.tab_stops.add_tab_stop(content_w, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t第 ")
    _add_field(hp.add_run(), "PAGE")
    hp.add_run(" 页 共 ")
    _add_field(hp.add_run(), "NUMPAGES")
    hp.add_run(" 页")

    # 正文：每行一段，每 lpp 行后强制分页（保证 >=lpp 行/页）
    n = len(selected)
    for i, ln in enumerate(selected):
        para = doc.add_paragraph()
        para.add_run(ln if ln != "" else " ")
        if (i + 1) % lpp == 0 and (i + 1) < n:
            para.add_run().add_break(WD_BREAK.PAGE)

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".", exist_ok=True)
    doc.save(out_path)


def _add_field(run, field):
    """在 run 上插入 Word 域（PAGE / NUMPAGES）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = f" {field} "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def main():
    ap = argparse.ArgumentParser(description="软著源代码 60 页鉴别材料生成器")
    ap.add_argument("--src", action="append", required=True, help="代码文件夹（可多次指定）")
    ap.add_argument("--name", default="", help="软件全称（页眉用，应与申请表一致）")
    ap.add_argument("--version", default="V1.0", help="版本号（默认 V1.0）")
    ap.add_argument("--out", default="源代码.docx", help="输出 docx 路径")
    ap.add_argument("--lines-per-page", type=int, default=50, help="每页行数（默认 50，>=50 合规）")
    ap.add_argument("--pages-each-side", type=int, default=30, help="前/后各取多少页（默认 30）")
    ap.add_argument("--ext", action="append", default=None, help="覆盖默认扩展名白名单（可多次，如 --ext .py）")
    ap.add_argument("--exclude-dir", action="append", default=[], help="额外排除目录名（可多次）")
    ap.add_argument("--include-tests", action="store_true", help="保留测试目录（默认排除 test/tests/spec/e2e 等）")
    ap.add_argument("--include-generated", action="store_true", help="保留自动生成代码（默认排除 moc_/ui_/*_pb2.py/*.designer.cs 等）")
    ap.add_argument("--tail-file", default=None, help="将路径含该子串的文件强制排到末尾，使末页落在程序结尾")
    ap.add_argument("--strip-comments", action="store_true", help="额外去掉纯括号行/极短注释行")
    ap.add_argument("--no-file-marker", action="store_true", help="不在每个文件前插入 // ===== 路径 ===== 标记行")
    ap.add_argument("--font", default="Consolas", help="等宽字体（默认 Consolas）")
    ap.add_argument("--font-size", type=float, default=9.0, help="字号 pt（默认 9）")
    ap.add_argument("--dry-run", action="store_true", help="只统计与自检，不生成 docx")
    args = ap.parse_args()

    exts = args.ext if args.ext else DEFAULT_EXTS
    files, excluded, stats = collect_files(
        args.src, exts, set(args.exclude_dir),
        include_tests=args.include_tests, include_generated=args.include_generated)
    if not files:
        print("[错误] 未收集到任何源码文件。请检查 --src 路径与 --ext 白名单。", file=sys.stderr)
        sys.exit(2)
    files = order_files(files, tail_pat=args.tail_file)
    root_marker = os.path.abspath(args.src[0]) if len(args.src) == 1 else None
    lines, total_lines, per_file, last_line = read_code_lines(
        files, root_marker, args.strip_comments, not args.no_file_marker)
    selected, pages, truncated = select_pages(lines, args.lines_per_page, args.pages_each_side)

    # 语言分布（按扩展名统计清洗后行数），便于发现多语言混排
    lang = {}
    for p, c in per_file:
        lang[os.path.splitext(p)[1].lower()] = lang.get(os.path.splitext(p)[1].lower(), 0) + c
    lang_sorted = sorted(lang.items(), key=lambda kv: -kv[1])

    # 统计与自检报告（始终打印）
    print("==== 软著源代码材料 统计 ====")
    print(f"收集文件数      : {len(files)}")
    print(f"语言分布(行)    : {', '.join(f'{e}={n}' for e, n in lang_sorted) or '无'}")
    print(f"排除目录(命中)  : {excluded if excluded else '无'}"
          + ("" if args.include_tests else "（含默认排除的测试目录）"))
    print(f"排除自动生成码  : {stats['generated_skipped']} 个文件"
          + ("（已用 --include-generated 保留）" if args.include_generated else ""))
    print(f"清洗后总代码行数: {total_lines} 行   <-- 回填申请表'源程序量'（记得带'行'字）")
    print(f"自然总页数      : {pages if not truncated else math.ceil(total_lines/args.lines_per_page)}（每页{args.lines_per_page}行）")
    if truncated:
        print(f"裁剪策略        : 超过 {2*args.pages_each_side} 页 -> 取前 {args.pages_each_side} 页 + 后 {args.pages_each_side} 页 = {pages} 页")
    else:
        print(f"裁剪策略        : 不足 {2*args.pages_each_side} 页 -> 全部提交（{pages} 页）")
    print(f"末文件          : {os.path.basename(files[-1]) if files else '无'}")
    print(f"末行(应为程序结尾): {last_line[:80]!r}")
    print("自检提醒        : 名称三处需一致(申请表/源码页眉/说明书页眉)；确认末页落在程序自然结尾。")
    # 多语言混排提醒
    big = [e for e, n in lang_sorted if n >= max(50, int(0.1 * total_lines))]
    if len(big) > 1:
        print(f"[提醒] 检测到多语言混排（{', '.join(big)}）。如只想收主语言，用 --ext 限定，例如 --ext .py。", file=sys.stderr)
    # 末页结尾提醒
    if truncated and not args.tail_file:
        print("[提醒] 已截断为前30+后30页，末页可能停在边角文件中段。"
              "可用 --tail-file <主程序文件名子串> 把它排到末尾，使末页落在程序结尾。", file=sys.stderr)

    if args.dry_run:
        print("[dry-run] 未生成 docx。")
        return
    if not args.name:
        print("[警告] 未提供 --name，页眉软件名为空；建议补上以与申请表一致。", file=sys.stderr)
    try:
        build_docx(selected, args.lines_per_page, args.name, args.version, args.out, args.font, args.font_size)
    except ImportError:
        print("[错误] 缺少 python-docx。请在 conda 环境 patent-drafting 中运行，或 pip install python-docx。", file=sys.stderr)
        sys.exit(3)
    print(f"[完成] 已生成: {os.path.abspath(args.out)}  （{pages} 页）")


if __name__ == "__main__":
    main()
