#!/usr/bin/env python
"""专利说明书附图生成器（patent-drafting）。

职责：只做 配置解析 → 严格渲染（mermaid_render --patent，白底黑线/失败即报错）→ A4 组版 → DOCX。
不做内容判断、不做验收（验收交 patent_figures_check.py）。

用法：
    python patent_figures.py --config figures.json --out-docx 说明书附图.docx [--assets-dir 附图] [--mmdc-scale 3]

配置（JSON）顶层：paper(A4)、margins_mm{top,bottom,left,right}、figures[...]。
每个 figure：number, subfigure(null|a|b), type, abstraction, source({mmd:路径}|{image:路径}),
full_name, is_abstract(bool), orientation(portrait|landscape), line_legend[], spec_section,
same_page_group(null|id)。长 mermaid 放各自 .mmd，不塞 JSON。
附图页只排居中“图N(±a/b)”，完整图名仅进说明书“附图说明”。
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

CN_NUM = "〇一二三四五六七八九十"
TARGET_DPI = 180  # 嵌入时有效分辨率下限：小图不过度放大（物理变小但清晰）


def _cn_index(n: int) -> str:
    if 1 <= n <= 10:
        return CN_NUM[n]
    return str(n)


def _strip_png_metadata(src: Path, dst: Path) -> tuple[int, int]:
    """重存 PNG 去除元数据（中性化），返回像素 (宽, 高)。"""
    from PIL import Image

    with Image.open(src) as im:
        im = im.convert("RGB") if im.mode not in ("RGB", "L") else im
        w, h = im.size
        clean = Image.new(im.mode, im.size)
        clean.paste(im)
        dst.parent.mkdir(parents=True, exist_ok=True)
        clean.save(dst, "PNG")  # 不带 pnginfo → 去文本块/时间戳
    return w, h


def _render_mmd_figures(figures: list[dict], cfg_dir: Path, assets_dir: Path, scale: float) -> dict:
    """把所有 mmd 源附图合并为一个 md，调 mermaid_render --patent 渲染，返回 {fig_key: png_path}。"""
    import mermaid_render

    mmd_figs = [f for f in figures if "mmd" in (f.get("source") or {})]
    out: dict = {}
    if not mmd_figs:
        return out

    blocks = []
    for f in mmd_figs:
        mmd_path = (cfg_dir / f["source"]["mmd"]).resolve()
        if not mmd_path.is_file():
            raise FileNotFoundError(f"找不到 mmd：{mmd_path}")
        blocks.append("```mermaid\n" + mmd_path.read_text(encoding="utf-8").rstrip() + "\n```\n")
    combined = "\n".join(blocks)

    work = assets_dir / "_render"
    work.mkdir(parents=True, exist_ok=True)
    tmp_md = work / "_combined.md"

    mermaid_render._PATENT_MMDC_CONFIG = mermaid_render._write_patent_mmdc_config()
    try:
        new_md, ok, fail = mermaid_render.render_markdown_mermaid(
            combined,
            out_md_path=tmp_md,
            assets_rel=".",
            mmdc_scale=scale,
            mmdc_width=1600,
            mmdc_height=1200,
        )
    finally:
        try:
            mermaid_render._PATENT_MMDC_CONFIG.unlink(missing_ok=True)
        except OSError:
            pass
    if fail:
        raise RuntimeError(f"--patent 严格渲染：{fail} 处 mermaid 失败，已硬失败（不降级）")
    if ok != len(mmd_figs):
        raise RuntimeError(f"渲染数({ok})与 mmd 附图数({len(mmd_figs)})不符")

    for i, f in enumerate(mmd_figs, start=1):
        src_png = work / f"fig_{i:03d}.png"
        key = _fig_key(f)
        dst_png = assets_dir / f"fig_{key}.png"
        _strip_png_metadata(src_png, dst_png)
        out[key] = dst_png
    return out


def _fig_key(f: dict) -> str:
    sub = f.get("subfigure")
    return f"{f['number']}{sub}" if sub else f"{f['number']}"


def _fig_caption(f: dict) -> str:
    sub = f.get("subfigure")
    return f"图{f['number']}({sub})" if sub else f"图{f['number']}"


def _add_page_number_footer(section) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    for t, attr in (("begin", "w:fldCharType"), (None, None), ("end", "w:fldCharType")):
        if t in ("begin", "end"):
            fld = OxmlElement("w:fldChar")
            fld.set(qn(attr), t)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = "PAGE"
            run._r.append(instr)


def build_docx(config: dict, cfg_dir: Path, assets_dir: Path, out_docx: Path, scale: float) -> dict:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Mm, Pt

    figures = config["figures"]
    margins = config.get("margins_mm", {"top": 25, "bottom": 15, "left": 25, "right": 15})

    rendered = _render_mmd_figures(figures, cfg_dir, assets_dir, scale)

    doc = Document()
    # 中性文档元数据
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.title = ""
    cp.comments = ""

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(margins["top"])
    sec.bottom_margin = Mm(margins["bottom"])
    sec.left_margin = Mm(margins["left"])
    sec.right_margin = Mm(margins["right"])
    _add_page_number_footer(sec)

    content_w_mm = 210 - margins["left"] - margins["right"]
    content_h_mm = 297 - margins["top"] - margins["bottom"] - 12  # 留题注/页脚

    n = len(figures)
    for idx, f in enumerate(figures):
        key = _fig_key(f)
        src = f.get("source") or {}
        if "mmd" in src:
            png = rendered[key]
        elif "image" in src:
            png = (cfg_dir / src["image"]).resolve()
            png = assets_dir / f"fig_{key}.png"
            _strip_png_metadata((cfg_dir / src["image"]).resolve(), png)
        else:
            raise ValueError(f"图{key} 缺 source.mmd 或 source.image")

        from PIL import Image

        with Image.open(png) as im:
            iw, ih = im.size
        # 适配版心；同时不把小图放大到有效 DPI 低于 TARGET_DPI（细长/简单图也清晰）
        max_w_by_dpi = iw / TARGET_DPI * 25.4
        target_w = min(content_w_mm, content_h_mm * iw / ih, max_w_by_dpi)

        # 图片段（居中，与题注保持同页）
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.keep_with_next = True
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(2)
        p_img.add_run().add_picture(str(png), width=Mm(target_w))

        # 题注段（仅“图N(±a/b)”，居中）
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(6)
        run = p_cap.add_run(_fig_caption(f))
        run.font.size = Pt(10.5)

        # 分页：默认一图一页；same_page_group 相同的相邻图不分页
        last = idx == n - 1
        if not last:
            cur_grp = f.get("same_page_group")
            nxt_grp = figures[idx + 1].get("same_page_group")
            same = cur_grp is not None and cur_grp == nxt_grp
            if not same:
                p_cap.add_run().add_break(WD_BREAK.PAGE)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    return {"figures": n, "assets": str(assets_dir)}


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="专利说明书附图生成器：配置→严格渲染→A4 DOCX")
    ap.add_argument("--config", required=True, type=Path, help="附图配置 JSON")
    ap.add_argument("--out-docx", required=True, type=Path, help="输出 A4 附图 .docx")
    ap.add_argument("--assets-dir", type=Path, default=None, help="PNG 资产目录（默认 out-docx 同级 figN_assets）")
    ap.add_argument("--mmdc-scale", type=float, default=3.0, help="mmdc 渲染缩放（默认 3，清晰）")
    args = ap.parse_args(argv)

    cfg_path = args.config.resolve()
    if not cfg_path.is_file():
        print(f"错误：找不到配置 {cfg_path}", file=sys.stderr)
        return 1
    config = json.loads(cfg_path.read_text(encoding="utf-8"))
    cfg_dir = cfg_path.parent
    out_docx = args.out_docx.resolve()
    assets_dir = (args.assets_dir.resolve() if args.assets_dir else out_docx.parent / "figs_assets")

    try:
        info = build_docx(config, cfg_dir, assets_dir, out_docx, args.mmdc_scale)
    except Exception as e:
        print(f"错误：{e}", file=sys.stderr)
        return 2
    print(f"已写入 A4 附图：{out_docx}（{info['figures']} 图，PNG 于 {info['assets']}）")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
